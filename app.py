import io
import os
from datetime import datetime
from typing import List

import streamlit as st
from docx import Document
from openai import OpenAI
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


APP_TITLE = "Neurospicy Translation"
APP_SUBTITLE = "Say it in your own words. Receive the format the other system expects."

TRAITS = [
    "Autism",
    "ADHD",
    "AuDHD",
    "Hyperlexia",
    "Dyslexia",
    "Dyspraxia",
    "Alexithymia",
    "Word-finding difficulty",
    "English is not my first language",
    "Prefer speaking",
    "Prefer typing",
    "Undiagnosed / exploring",
]

OUTPUT_TYPES = [
    "General clear English",
    "NHS or GP communication",
    "PIP or benefits form",
    "Local government or council",
    "Housing",
    "Employer or reasonable adjustments",
    "Complaint",
    "Appeal",
    "Relationship or personal message",
    "Evidence summary",
    "Action plan",
]

OUTPUT_STYLES = [
    "Clear and neutral",
    "Warm and collaborative",
    "Professional",
    "Formal",
    "Very concise",
    "Detailed",
    "Bullet points",
    "Easy Read",
]

st.set_page_config(
    page_title=APP_TITLE,
    page_icon="🧠",
    layout="centered",
)

st.markdown(
    """
    <style>
    .stApp {
        background: #F4F1E8;
        color: #28352F;
    }
    .block-container {
        max-width: 900px;
        padding-top: 2rem;
        padding-bottom: 4rem;
    }
    h1, h2, h3 {
        color: #354A42;
    }
    div[data-baseweb="select"] > div,
    textarea,
    input {
        background-color: #FCFBF6 !important;
    }
    .quiet-box {
        padding: 1rem;
        border-radius: 0.75rem;
        background: #E8ECE7;
        border: 1px solid #C8D0C9;
        margin-bottom: 1rem;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

st.title(APP_TITLE)
st.caption(APP_SUBTITLE)

st.markdown(
    """
    <div class="quiet-box">
    Your communication is not defective. Many services require information in a narrow,
    linear format. This tool uses AI to reorganise your meaning without treating your
    natural language as a problem.
    </div>
    """,
    unsafe_allow_html=True,
)

with st.expander("Privacy and limits", expanded=False):
    st.write(
        "This starter version does not create user accounts or save translations in a database. "
        "Text is sent to the configured AI provider to generate the translation. "
        "Do not paste passwords, bank details, or unnecessary identifying information. "
        "AI output can be wrong, especially for legal, medical, or benefits matters. "
        "Check the final text before sending it."
    )

st.subheader("1. Tell us how your communication system works")
traits: List[str] = st.multiselect(
    "Choose any traits or preferences that help the translator interpret your words",
    TRAITS,
    placeholder="Select any that fit",
)

other_traits = st.text_input(
    "Anything else? Optional",
    placeholder="For example: I think in maps, I lose nouns, or I need plain language.",
)

st.subheader("2. Say it in your own words")
natural_input = st.text_area(
    "No need to organise it first",
    height=220,
    placeholder=(
        "Type naturally. Tangents, fragments, repetition, context and emotion are allowed. "
        "The translator will map the information before rewriting it."
    ),
)

st.subheader("3. What form does the outside world expect?")
output_type = st.selectbox("Purpose", OUTPUT_TYPES)
output_style = st.selectbox("Style", OUTPUT_STYLES)

context = st.text_area(
    "Extra context or instructions, optional",
    height=120,
    placeholder=(
        "Who will read it? What outcome do you need? Is there a word limit, deadline, "
        "form question, or earlier correspondence?"
    ),
)

preserve_emotion = st.checkbox(
    "Preserve emotional meaning where relevant",
    value=True,
)
include_questions = st.checkbox(
    "Ask clarifying questions instead of guessing when important facts are missing",
    value=True,
)

SYSTEM_PROMPT = """
You are an accessibility-focused communication translator.

Your job is to translate natural, networked, associative, gestalt, nonlinear, dyslexic,
hyperlexic, autistic, ADHD, alexithymic, multilingual, or otherwise non-standard input into
the format expected by the selected audience.

Core rules:
1. Assume competence.
2. Treat word-finding difficulty, fragments, repetition, tangents and relational descriptions
   as information-routing differences, not lack of understanding.
3. Preserve the user's intent, facts, boundaries, agency and desired outcome.
4. Do not add facts, diagnoses, legal claims or events that the user did not provide.
5. Separate fact, interpretation, emotional impact and requested action where useful.
6. For institutional outputs, organise information into a clear linear structure.
7. Use direct language without making the user sound submissive, childish or excessively polite.
8. If essential information is missing, mark it clearly or ask concise clarifying questions.
9. Do not claim legal, medical or benefits certainty.
10. Return only the translated output followed, when needed, by a short section titled
    "Questions before sending".
"""

def build_prompt() -> str:
    selected_traits = ", ".join(traits) if traits else "No traits selected"
    emotion_instruction = (
        "Preserve relevant emotional impact."
        if preserve_emotion
        else "Reduce emotional language while preserving the underlying facts and impact."
    )
    uncertainty_instruction = (
        "Ask short clarifying questions rather than guessing."
        if include_questions
        else "Use clearly marked placeholders for missing information and never guess."
    )

    return f"""
USER COMMUNICATION PROFILE
Selected traits/preferences: {selected_traits}
Additional description: {other_traits or "None provided"}

TARGET
Purpose: {output_type}
Style: {output_style}
Additional context: {context or "None provided"}

TRANSLATION INSTRUCTIONS
{emotion_instruction}
{uncertainty_instruction}

NATURAL INPUT
{natural_input}
""".strip()


def make_docx(text: str) -> bytes:
    document = Document()
    document.add_heading("Neurospicy Translation", level=1)
    document.add_paragraph(f"Created: {datetime.now().strftime('%d %B %Y, %H:%M')}")
    document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_pdf(text: str) -> bytes:
    buffer = io.BytesIO()
    styles = getSampleStyleSheet()
    story = [
        Paragraph("Neurospicy Translation", styles["Title"]),
        Spacer(1, 12),
        Paragraph(
            f"Created: {datetime.now().strftime('%d %B %Y, %H:%M')}",
            styles["Normal"],
        ),
        Spacer(1, 16),
    ]
    for paragraph in text.split("\n"):
        if paragraph.strip():
            safe = (
                paragraph.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            story.append(Paragraph(safe, styles["BodyText"]))
            story.append(Spacer(1, 8))
    SimpleDocTemplate(buffer, pagesize=A4).build(story)
    return buffer.getvalue()


if "translation" not in st.session_state:
    st.session_state.translation = ""

translate_clicked = st.button(
    "Map and translate",
    type="primary",
    use_container_width=True,
)

if translate_clicked:
    if not natural_input.strip():
        st.error("Please add something in your own words first.")
    elif not os.getenv("OPENAI_API_KEY"):
        st.error(
            "The server is missing OPENAI_API_KEY. Add it to your environment or Streamlit secrets."
        )
    else:
        try:
            with st.spinner("Mapping your ideas..."):
                client = OpenAI()
                response = client.responses.create(
                    model=os.getenv("OPENAI_MODEL", "gpt-5-mini"),
                    instructions=SYSTEM_PROMPT,
                    input=build_prompt(),
                )
                st.session_state.translation = response.output_text.strip()
        except Exception as exc:
            st.error(f"Translation failed: {exc}")

if st.session_state.translation:
    st.subheader("Translation")
    edited_translation = st.text_area(
        "Review and edit before using",
        value=st.session_state.translation,
        height=320,
    )

    st.download_button(
        "Download Word document",
        data=make_docx(edited_translation),
        file_name="neurospicy_translation.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

    st.download_button(
        "Download PDF",
        data=make_pdf(edited_translation),
        file_name="neurospicy_translation.pdf",
        mime="application/pdf",
        use_container_width=True,
    )

    st.download_button(
        "Download plain text",
        data=edited_translation.encode("utf-8"),
        file_name="neurospicy_translation.txt",
        mime="text/plain",
        use_container_width=True,
    )

    if st.button("Clear this translation", use_container_width=True):
        st.session_state.translation = ""
        st.rerun()

st.divider()
st.caption(
    "Prototype only. Before public launch, add consent wording, safeguarding routes, "
    "usage limits, payments, accessibility testing and a proper privacy policy."
)
